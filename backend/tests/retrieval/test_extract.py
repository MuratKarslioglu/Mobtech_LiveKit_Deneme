import io

import docx
import pytest
from pypdf import PdfWriter

from app.retrieval.extract import UnsupportedDocumentType, extract_text


def test_extract_text_from_plain_text():
    data = "merhaba dünya".encode("utf-8")
    assert extract_text(data, "text/plain") == "merhaba dünya"


def test_extract_text_from_docx():
    document = docx.Document()
    document.add_paragraph("birinci paragraf")
    document.add_paragraph("ikinci paragraf")
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text(
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "birinci paragraf" in text
    assert "ikinci paragraf" in text


def test_extract_text_from_pdf():
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)

    # Boş bir sayfanın metni boş string olur ama tür destekleniyor olmalı ve
    # hata fırlatmamalı.
    text = extract_text(buffer.getvalue(), "application/pdf")
    assert text == ""


def test_extract_text_raises_for_unsupported_type():
    with pytest.raises(UnsupportedDocumentType):
        extract_text(b"data", "application/zip")
